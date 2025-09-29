import streamlit as st
from agent import evaluate_use_case
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Streamlit page config
st.set_page_config(page_title="AI Ethics Agent", layout="centered")
st.title("🛡️ AI Ethics Agent")
st.markdown("""
Use this tool to **evaluate your AI use case** across key ethical risk categories:
- **Bias**
- **Privacy**
- **Explainability**
- **Governance & Oversight**

💡 Understanding these areas is critical before deploying any AI solution — especially for systems that impact real people. This tool helps uncover ethical risks you may not have considered.
  
_One report is generated per use case and downloadable as a Word document (.docx)._""")

# Input box
user_input = st.text_area(
    "Describe your AI use case:",
    placeholder="E.g. An AI model that predicts hospital readmission using patient records, deployed on cloud...",
    height=200,
    key="use_case_input"
)

if st.button("Evaluate Use Case"):
    if not user_input.strip():
        st.warning("Please enter a description before submitting.")
    else:
        progress = st.progress(0, text="Evaluating ethical risks...")

        results = evaluate_use_case(user_input)

        progress.progress(100, text="✅ Evaluation complete!")

        # Build Markdown-style string for report
        markdown_content = f"# AI Use Case Evaluation Report\n\n**Use Case:**\n{user_input.strip()}\n\n"
        for key, value in results.items():
            markdown_content += f"## {key}\n{value.strip()}\n\n"

        # Save Markdown file to reports/
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        reports_dir = Path("reports")
        reports_dir.mkdir(exist_ok=True)
        md_path = reports_dir / f"eval_{timestamp}.md"
        md_path.write_text(markdown_content, encoding="utf-8")

        # Generate .docx from markdown_content (no saving to disk)
        doc = Document()
        styles = doc.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(11)

        for line in markdown_content.strip().split("\n"):
            line = line.strip()
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("**") and line.endswith("**"):
                para = doc.add_paragraph()
                run = para.add_run(line[2:-2])
                run.bold = True
            elif "**" in line:
                para = doc.add_paragraph()
                segments = line.split("**")
                for i, segment in enumerate(segments):
                    run = para.add_run(segment)
                    if i % 2 == 1:
                        run.bold = True
            elif line.startswith("-") or line.startswith("\u2022"):
                para = doc.add_paragraph(line[1:].strip(), style="List Bullet")
            elif line:
                doc.add_paragraph(line)

        # Display report
        st.subheader("✅ Evaluation Results")
        st.markdown(markdown_content)

        # Download .docx (in-memory only)
        from io import BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        st.download_button(
            "📄 Download Word Report (.docx)",
            data=docx_buffer,
            file_name=f"eval_{timestamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
